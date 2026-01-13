from docx import Document

doc = Document(r'd:\devnolife\network\Modul_CW6552021557_Praktikum_Advanced_Network_Security_and_Protocols.docx')

# Extract all paragraphs
content = []
for para in doc.paragraphs:
    if para.text.strip():
        content.append(para.text)

# Also extract tables
for table in doc.tables:
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            if cell.text.strip():
                row_text.append(cell.text.strip())
        if row_text:
            content.append(' | '.join(row_text))

# Save to file
with open(r'd:\devnolife\network\modul_praktikum.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(content))

print("Extracted successfully to modul_praktikum.txt")
print("\n" + "="*50)
print("CONTENT:")
print("="*50 + "\n")
print('\n'.join(content))
